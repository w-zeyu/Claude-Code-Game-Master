#!/usr/bin/env python3
"""
Content extractors for various document formats.
Extract plain text from PDFs, Word documents, Markdown, and text files.
"""

import os
import re
from pathlib import Path
from typing import Optional


class PDFExtractor:
    """Extract text content from PDF files."""

    def __init__(self):
        """Initialize PDF extractor."""
        self.pypdf_available = False
        self.pdfplumber_available = False

        try:
            import PyPDF2
            self.pypdf_available = True
            self.PyPDF2 = PyPDF2
        except ImportError:
            pass

        try:
            import pdfplumber
            self.pdfplumber_available = True
            self.pdfplumber = pdfplumber
        except ImportError:
            pass

        if not self.pypdf_available and not self.pdfplumber_available:
            print("Warning: No PDF libraries available. Install PyPDF2 or pdfplumber.")

    def extract(self, filepath: str) -> str:
        """
        Extract text from a PDF file.

        Args:
            filepath: Path to the PDF file

        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"PDF file not found: {filepath}")

        # Try pdfplumber first (better for complex layouts)
        if self.pdfplumber_available:
            try:
                return self._extract_with_pdfplumber(filepath)
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
                if self.pypdf_available:
                    print("Falling back to PyPDF2...")

        # Fall back to PyPDF2
        if self.pypdf_available:
            try:
                return self._extract_with_pypdf2(filepath)
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")

        raise RuntimeError("No PDF extraction library available. Install PyPDF2 or pdfplumber.")

    def _extract_with_pdfplumber(self, filepath: str) -> str:
        """Extract text using pdfplumber."""
        text = []
        with self.pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- Page {page_num} ---\n")
                        text.append(page_text)
                        text.append("\n\n")
                except Exception as e:
                    print(f"Error extracting page {page_num}: {e}")

        return ''.join(text)

    def _extract_with_pypdf2(self, filepath: str) -> str:
        """Extract text using PyPDF2."""
        text = []
        with open(filepath, 'rb') as file:
            pdf_reader = self.PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- Page {page_num + 1} ---\n")
                        text.append(page_text)
                        text.append("\n\n")
                except Exception as e:
                    print(f"Error extracting page {page_num + 1}: {e}")

        return ''.join(text)


class MarkdownExtractor:
    """Extract text content from Markdown files."""

    def extract(self, filepath: str) -> str:
        """
        Extract text from a Markdown file.

        Args:
            filepath: Path to the Markdown file

        Returns:
            Text content with basic Markdown formatting preserved
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Markdown file not found: {filepath}")

        with open(filepath, 'r', encoding='utf-8') as file:
            content = file.read()

        # Clean up excessive whitespace while preserving structure
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content


class DocxExtractor:
    """Extract text content from Word documents."""

    def __init__(self):
        """Initialize DOCX extractor."""
        self.docx_available = False

        try:
            import docx
            self.docx_available = True
            self.docx = docx
        except ImportError:
            print("Warning: python-docx not available. Install python-docx to process Word documents.")

    def extract(self, filepath: str) -> str:
        """
        Extract text from a Word document.

        Args:
            filepath: Path to the DOCX file

        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Word document not found: {filepath}")

        if not self.docx_available:
            # Try basic text extraction as fallback
            return self._basic_extract(filepath)

        try:
            doc = self.docx.Document(filepath)
            text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
                    text.append('\n')

            # Extract tables
            for table in doc.tables:
                text.append('\n--- Table ---\n')
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text.append(' | '.join(row_text))
                    text.append('\n')
                text.append('\n')

            return ''.join(text)

        except Exception as e:
            print(f"Error extracting Word document: {e}")
            return self._basic_extract(filepath)

    def _basic_extract(self, filepath: str) -> str:
        """Basic text extraction fallback for Word documents."""
        # This is a very basic fallback that may not work well
        try:
            with open(filepath, 'rb') as file:
                content = file.read()
                # Try to decode as UTF-8, ignoring errors
                text = content.decode('utf-8', errors='ignore')
                # Remove binary garbage
                text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                return text
        except Exception as e:
            raise RuntimeError(f"Cannot extract text from Word document: {e}")


class TextExtractor:
    """Extract content from plain text files."""

    def extract(self, filepath: str) -> str:
        """
        Extract text from a plain text file.

        Args:
            filepath: Path to the text file

        Returns:
            File content
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Text file not found: {filepath}")

        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue

        # If all encodings fail, try with error handling
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()


def extract_content(filepath: str) -> str:
    """
    Extract text content from any supported file type.

    Args:
        filepath: Path to the file

    Returns:
        Extracted text content
    """
    file_path = Path(filepath)
    extension = file_path.suffix.lower()

    extractors = {
        '.pdf': PDFExtractor(),
        '.md': MarkdownExtractor(),
        '.markdown': MarkdownExtractor(),
        '.docx': DocxExtractor(),
        '.doc': DocxExtractor(),
        '.txt': TextExtractor(),
        '.text': TextExtractor(),
    }

    extractor = extractors.get(extension)
    if not extractor:
        raise ValueError(f"Unsupported file type: {extension}")

    return extractor.extract(filepath)


class ContentExtractor:
    """Unified content extractor for any supported file type."""

    def extract_text(self, filepath: str) -> str:
        """Extract text from any supported file type."""
        return extract_content(filepath)


def main():
    """Test the extractors."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: content_extractor.py <file>")
        sys.exit(1)

    filepath = sys.argv[1]

    try:
        content = extract_content(filepath)
        print(f"Extracted {len(content)} characters from {filepath}")
        print("\nFirst 500 characters:")
        print("-" * 50)
        print(content[:500])
        print("-" * 50)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()